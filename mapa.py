from docx import Document
from docx.shared import Inches

# Crear el documento Word
doc = Document()

# Título
doc.add_heading('Análisis FODA y Mapa Conceptual del Departamento de Desarrollo', level=1)

# Datos generales
doc.add_paragraph('Nombre del estudiante: Miguel')
doc.add_paragraph('Carrera: Desarrollo de Software')
doc.add_paragraph('Fecha de entrega: [Completa la fecha]')
doc.add_paragraph('Facilitadora: Lic. Rocío del Carmen Durán Maldonado')
doc.add_paragraph('Nombre de la empresa: SoftLabs Solutions S.A. de C.V.')
doc.add_paragraph('Área analizada: Departamento de Desarrollo de Software')

# Secciones
doc.add_heading('1. Introducción', level=2)
doc.add_paragraph(
    "El análisis FODA es una herramienta fundamental dentro del proceso de planeación estratégica. "
    "En el área de desarrollo de software, su aplicación permite identificar las fortalezas y debilidades internas, "
    "así como las oportunidades y amenazas del entorno, con el objetivo de formular propuestas de mejora que impulsen "
    "el rendimiento del equipo y el cumplimiento de los objetivos de la empresa."
)

doc.add_heading('2. Información general de la empresa', level=2)
doc.add_paragraph(
    "Nombre: SoftLabs Solutions S.A. de C.V.\n"
    "Giro: Desarrollo de software a la medida para MIPYMES.\n"
    "Misión: Crear soluciones tecnológicas personalizadas que optimicen los procesos de las MIPYMES.\n"
    "Visión: Ser reconocida a nivel nacional por su innovación, calidad y compromiso con el desarrollo tecnológico.\n"
    "Valores: Innovación, trabajo en equipo, compromiso con el cliente, ética profesional, mejora continua."
)

doc.add_heading('3. Análisis FODA del Departamento de Desarrollo', level=2)

# Tabla FODA
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fortalezas'
hdr_cells[1].text = 'Oportunidades'

foda_data = [
    ("- Equipo con habilidades en desarrollo web y móvil", "- Aumento de la demanda de software personalizado"),
    ("- Uso de metodologías ágiles", "- Alianzas con universidades y empresas tecnológicas"),
    ("- Cultura de aprendizaje continuo", "- Acceso a tecnologías emergentes (IA, nube, blockchain)"),
    ("Debilidades", "Amenazas"),
    ("- Dificultades en la gestión del tiempo", "- Competencia creciente de freelancers y agencias grandes"),
    ("- Falta de experiencia en nuevas tecnologías", "- Cambios constantes en las necesidades del mercado"),
    ("- Documentación incompleta o ineficiente", "- Riesgos de ciberseguridad")
]

for row_data in foda_data:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

doc.add_heading('4. Diagnóstico del área', level=2)
doc.add_paragraph(
    "El área de desarrollo cuenta con fortalezas significativas como su dominio técnico en plataformas populares "
    "y su capacidad de adaptación a metodologías ágiles. Sin embargo, enfrenta debilidades importantes en la planificación "
    "de tiempos y en la actualización constante ante nuevas tecnologías. Las oportunidades externas representan una vía clara "
    "de crecimiento y posicionamiento, mientras que las amenazas exigen estar preparados ante cambios rápidos en el entorno tecnológico."
)

doc.add_heading('5. Propuestas de mejora', level=2)
doc.add_paragraph(
    "1. Impulsar la capacitación técnica en nuevas tecnologías (IA, blockchain, etc.).\n"
    "2. Adoptar herramientas de gestión de proyectos como Jira o ClickUp.\n"
    "3. Estandarizar y mejorar la documentación técnica interna.\n"
    "4. Establecer convenios con universidades para integrar talento nuevo.\n"
    "5. Diseñar e implementar un plan de ciberseguridad y contingencias."
)

doc.add_heading('6. Mapa Conceptual', level=2)
doc.add_paragraph('[Inserta aquí el mapa conceptual visual creado en Canva, CmapTools, PowerPoint, etc.]')

doc.add_heading('7. Conclusión', level=2)
doc.add_paragraph(
    "La aplicación del análisis FODA permite comprender de manera objetiva la situación actual del área de desarrollo, "
    "identificando tanto sus ventajas como sus retos. A partir de ello, es posible establecer estrategias claras y realistas "
    "que contribuyan a mejorar el desempeño del equipo, mantenerse competitivo en el mercado y adaptarse a las exigencias del entorno tecnológico."
)

doc.add_heading('8. Bibliografía', level=2)
doc.add_paragraph(
    "- Ramírez Rojas, J.L. (2009). Procedimiento para la elaboración de un análisis FODA. Universidad Veracruzana.\n"
    "- Kaplan, R. & Norton, D. (2001). El Balanced Scorecard: Cuadro de Mando Integral.\n"
    "- Álvarez, M.T., Moreno, S.A. y Chávez, M.Y. (2020). El BSC como herramienta estratégica."
)

# Guardar el documento
doc_path = "/mnt/data/Actividad_FODA_Desarrollo_Software.docx"
doc.save(doc_path)

doc_path
